from pathlib import Path

from docx import Document
from openpyxl import Workbook

from brp.adapters.docs_manual import ManualDocumentAdapter
from brp.ir.canonical import canonical_bytes


def make_xlsx(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "가입규칙"
    sheet.append(
        [
            "decision_key",
            "decision_name",
            "rule_id",
            "input_name",
            "input_type",
            "operator",
            "operand_value",
            "output_name",
            "output_type",
            "output_value",
            "section",
        ]
    )
    sheet.append(
        [
            "manual_eligibility",
            "수동 가입 자격",
            "R001",
            "age",
            "integer",
            "LT",
            18,
            "reasonCode",
            "string",
            "미성년",
            "가입 나이",
        ]
    )
    workbook.save(path)


def test_xlsx_candidate_preserves_korean_sheet_section_and_cell(tmp_path: Path) -> None:
    path = tmp_path / "가입규칙.xlsx"
    make_xlsx(path)
    adapter = ManualDocumentAdapter([path])
    batch = adapter.extract(adapter.discover(None)[0])
    rule = batch.decisions[0].content.rules[0]
    reference = rule.source_references[0]
    assert reference.sheet == "가입규칙"
    assert reference.section == "가입 나이"
    assert reference.cell_range == "A2:K2"
    assert rule.confidence == 0.35
    assert "미성년" in canonical_bytes(batch.decisions[0].content).decode("utf-8")


def test_sparse_docx_text_enters_review_queue(tmp_path: Path) -> None:
    path = tmp_path / "manual.docx"
    document = Document()
    document.add_heading("가입 조건", level=1)
    document.add_paragraph("경우에 따라 추가 심사가 필요합니다.")
    document.save(path)
    adapter = ManualDocumentAdapter([path])
    batch = adapter.extract(adapter.discover(None)[0])
    assert not batch.decisions
    assert batch.unmappable[0].reason_code == "AMBIGUOUS_MANUAL_TEXT"
    assert batch.unmappable[0].provenance["section"] == "가입 조건"
